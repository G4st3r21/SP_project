from datetime import datetime
from openpyxl import Workbook
import zipfile
import docx
import os


# def has_images(path):
#     z = zipfile.ZipFile(path)
#
#     all_files = z.namelist()
#
#     images = [i for i in filter(lambda x: x.startswith('word/media/'), all_files)]
#     print('Количество изображений: ', len(images))
#
#     return len(images)


def is_digit(string):
    if string.isdigit():
        return True
    else:
        try:
            float(string)
            return True
        except ValueError:
            return False


def print_properties(doc, title):
    properties = doc.core_properties
    print('Название документа:', title)
    print('Автор документа:', properties.author)
    print('Автор последней правки:', properties.last_modified_by)
    print('Дата создания документа:', properties.created)
    print('Дата последней правки:', properties.modified)
    print('Дата последней печати:', properties.last_printed)
    print('Количество сохранений:', properties.revision)
    print('КОЛИЧЕСТВО ТАБЛИЦ:', len(doc.tables))
    print('P.s. Количество таблиц является не точным, т.к. в файле используются псевдотаблицы\n', )


def print_table_info(table, count):
    print('Таблица №' + str(count))
    print('Количество строк:', len(table.rows))
    print('Количество столбцов:', len(table.columns))
    if len(table.rows) == 1:
        print('Это ПСЕВДОТАБЛИЦА, которая не будет перенесена в excel')
    print()


def check_for_merged_cells(table):
    unique, merged = [], []
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            cell_loc = (tc.top, tc.bottom, tc.left, tc.right)
            if cell_loc in unique:
                merged.append((cell_loc, abs(cell_loc[0] - cell_loc[1]), abs(cell_loc[2] - cell_loc[3])))
            else:
                unique.append(cell_loc)

    return merged


def create_xlsx_table(table, merged, excel_file, count):
    if len(table.rows) > 1:
        excel_sheet = excel_file.create_sheet(index=count)

        for row_count, row in enumerate(table.rows, start=1):
            for cell_count, cell in enumerate(row.cells, start=1):
                try:
                    temp_value = cell.text.split(',')
                    if temp_value[1].isdigit() and temp_value[0].isdigit() and len(temp_value) == 2:
                        excel_sheet.cell(row_count, cell_count).value = float(temp_value[0] + '.' + temp_value[1])
                    else:
                        excel_sheet.cell(row_count, cell_count).value = str(cell.text)
                except IndexError:
                    if str(cell.text).isdigit():
                        excel_sheet.cell(row_count, cell_count).value = int(cell.text)
                    else:
                        excel_sheet.cell(row_count, cell_count).value = str(cell.text)

        for cell in merged:
            excel_sheet.merge_cells(start_row=int(cell[0][0] + 1), start_column=int(cell[0][2] + 1),
                                    end_row=int(cell[0][1]), end_column=int(cell[0][3]))


start_time = datetime.now()
paths = []
folder = 'wordFiles'
problem_files = []
files_without_tables = []
# files_with_images = []

for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))

for path in paths:

    doc = docx.Document(path)
    print_properties(doc, path.split('\\')[-1])

    # if has_images(path) > 50:
    #     files_with_images.append(path)

    tables = doc.tables

    excel_file = Workbook()
    for count, table in enumerate(tables, start=1):
        print_table_info(table, count)
        try:
            merged = check_for_merged_cells(table)
        except Exception:
            problem_files.append((path.split('\\')[-1].split('.')[0] + '.xlsx', 'Номер таблицы в файле: ' + str(count)))
            merged = []
        create_xlsx_table(table, merged, excel_file, count)

    del excel_file['Sheet']

    sheets = excel_file.sheetnames
    if sheets:
        print(sheets)
        excel_file.save(filename=path.split('\\')[-1].split('.')[0] + '.xlsx')
    else:
        files_without_tables.append(path)


print('Время выполнения:', datetime.now() - start_time)

print('Файлы, которые скрипт не смог обработать:')
print(*problem_files, sep='\n')

print('Файлы без таблиц:')
print(*files_without_tables, sep='\n')

# print(f'Файлы в которых присутствуют изображения(больше 50 изображений)({len(files_with_images)}):')
# print(*files_with_images, sep='\n')
